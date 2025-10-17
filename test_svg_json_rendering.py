#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""测试 SVG-in-JSON 格式解析和渲染"""

import json
import re
from io import BytesIO
from PIL import Image
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 用户提供的示例 LLM 输出
EXAMPLE_OUTPUT = """{
  "text": "1. (0分) (2016高二上重庆期中) 如图所示，在正方体 $ABCD-A_1B_1C_1D_1$ 中， $E$ 、 $F$ 分别是 $AB$ 和 $AA_1$ 的中点．证明：\\n(1) $E$ 、 $C$ 、 $D_1$ 、 $F$ 四点共面；\\n(2) $CE$ 、 $D_1F$ 、 $DA$ 三线共点．",
  "figure_svg": "<svg viewBox=\\"0 0 225 155\\" xmlns=\\"http://www.w3.org/2000/svg\\">\\n  <!-- 后面顶点 -->\\n  <line x1=\\"20\\" y1=\\"80\\" x2=\\"100\\" y2=\\"80\\" stroke=\\"#333\\" stroke-width=\\"1.5\\" stroke-dasharray=\\"3,3\\"/>\\n  <line x1=\\"100\\" y1=\\"80\\" x2=\\"100\\" y2=\\"20\\" stroke=\\"#333\\" stroke-width=\\"1.5\\"/>\\n  <line x1=\\"100\\" y1=\\"20\\" x2=\\"20\\" y2=\\"20\\" stroke=\\"#333\\" stroke-width=\\"1.5\\" stroke-dasharray=\\"3,3\\"/>\\n  <line x1=\\"20\\" y1=\\"20\\" x2=\\"20\\" y2=\\"80\\" stroke=\\"#333\\" stroke-width=\\"1.5\\" stroke-dasharray=\\"3,3\\"/>\\n  \\n  <!-- 前面顶点 -->\\n  <line x1=\\"80\\" y1=\\"140\\" x2=\\"160\\" y2=\\"140\\" stroke=\\"#333\\" stroke-width=\\"2\\"/>\\n  <line x1=\\"160\\" y1=\\"140\\" x2=\\"160\\" y2=\\"80\\" stroke=\\"#333\\" stroke-width=\\"2\\"/>\\n  <line x1=\\"160\\" y1=\\"80\\" x2=\\"80\\" y2=\\"80\\" stroke=\\"#333\\" stroke-width=\\"2\\"/>\\n  <line x1=\\"80\\" y1=\\"80\\" x2=\\"80\\" y2=\\"140\\" stroke=\\"#333\\" stroke-width=\\"2\\"/>\\n  \\n  <!-- 连接线 -->\\n  <line x1=\\"20\\" y1=\\"80\\" x2=\\"80\\" y2=\\"140\\" stroke=\\"#333\\" stroke-width=\\"1.5\\" stroke-dasharray=\\"3,3\\"/>\\n  <line x1=\\"100\\" y1=\\"80\\" x2=\\"160\\" y2=\\"140\\" stroke=\\"#333\\" stroke-width=\\"2\\"/>\\n  <line x1=\\"100\\" y1=\\"20\\" x2=\\"160\\" y2=\\"80\\" stroke=\\"#333\\" stroke-width=\\"2\\"/>\\n  <line x1=\\"20\\" y1=\\"20\\" x2=\\"80\\" y2=\\"80\\" stroke=\\"#333\\" stroke-width=\\"2\\"/>\\n  \\n  <!-- 标注点 -->\\n  <circle cx=\\"80\\" cy=\\"140\\" r=\\"2\\" fill=\\"#333\\"/>\\n  <circle cx=\\"160\\" cy=\\"140\\" r=\\"2\\" fill=\\"#333\\"/>\\n  <circle cx=\\"160\\" cy=\\"80\\" r=\\"2\\" fill=\\"#333\\"/>\\n  <circle cx=\\"80\\" cy=\\"80\\" r=\\"2\\" fill=\\"#333\\"/>\\n  <circle cx=\\"20\\" cy=\\"80\\" r=\\"2\\" fill=\\"#333\\"/>\\n  <circle cx=\\"100\\" cy=\\"80\\" r=\\"2\\" fill=\\"#333\\"/>\\n  <circle cx=\\"100\\" cy=\\"20\\" r=\\"2\\" fill=\\"#333\\"/>\\n  <circle cx=\\"20\\" cy=\\"20\\" r=\\"2\\" fill=\\"#333\\"/>\\n  \\n  <!-- E点（AB中点） -->\\n  <circle cx=\\"120\\" cy=\\"140\\" r=\\"2\\" fill=\\"red\\"/>\\n  <text x=\\"120\\" y=\\"155\\" font-size=\\"10\\" fill=\\"#333\\" text-anchor=\\"middle\\">E</text>\\n  \\n  <!-- F点（AA1中点） -->\\n  <circle cx=\\"50\\" cy=\\"50\\" r=\\"2\\" fill=\\"red\\"/>\\n  <text x=\\"50\\" y=\\"45\\" font-size=\\"10\\" fill=\\"#333\\" text-anchor=\\"middle\\">F</text>\\n  \\n  <!-- 顶点标注 -->\\n  <text x=\\"80\\" y=\\"155\\" font-size=\\"10\\" fill=\\"#333\\" text-anchor=\\"middle\\">A</text>\\n  <text x=\\"160\\" y=\\"155\\" font-size=\\"10\\" fill=\\"#333\\" text-anchor=\\"middle\\">B</text>\\n  <text x=\\"160\\" y=\\"75\\" font-size=\\"10\\" fill=\\"#333\\" text-anchor=\\"middle\\">C</text>\\n  <text x=\\"80\\" y=\\"75\\" font-size=\\"10\\" fill=\\"#333\\" text-anchor=\\"middle\\">D</text>\\n  <text x=\\"20\\" y=\\"95\\" font-size=\\"10\\" fill=\\"#333\\" text-anchor=\\"middle\\">A₁</text>\\n  <text x=\\"100\\" y=\\"95\\" font-size=\\"10\\" fill=\\"#333\\" text-anchor=\\"middle\\">B₁</text>\\n  <text x=\\"100\\" y=\\"15\\" font-size=\\"10\\" fill=\\"#333\\" text-anchor=\\"middle\\">C₁</text>\\n  <text x=\\"20\\" y=\\"15\\" font-size=\\"10\\" fill=\\"#333\\" text-anchor=\\"middle\\">D₁</text>\\n</svg>"
}"""


def parse_svg_json(content: str) -> tuple[str, str]:
    """
    解析包含 text 和 figure_svg 的 JSON 格式

    Returns:
        tuple: (text, svg_content)
    """
    try:
        data = json.loads(content)
        text = data.get('text', '')
        svg_content = data.get('figure_svg', '')
        return text, svg_content
    except json.JSONDecodeError as e:
        print(f"JSON 解析失败: {e}")
        return '', ''


def svg_to_png(svg_content: str, output_width: int = 800) -> Image.Image:
    """
    将 SVG 转换为 PNG 图像

    Args:
        svg_content: SVG 字符串
        output_width: 输出宽度（像素）

    Returns:
        PIL Image object
    """
    try:
        import cairosvg

        # 转换 SVG 到 PNG bytes
        png_bytes = cairosvg.svg2png(bytestring=svg_content.encode('utf-8'),
                                      output_width=output_width)

        # 转换为 PIL Image
        image = Image.open(BytesIO(png_bytes))
        return image

    except ImportError:
        print("cairosvg 未安装，请运行: uv pip install cairosvg")
        raise
    except Exception as e:
        print(f"SVG 转换失败: {e}")
        raise


def create_word_document(text: str, image: Image.Image, output_path: str):
    """
    创建包含文本和图片的 Word 文档

    Args:
        text: 题目文字
        image: PIL Image 对象
        output_path: 输出文件路径
    """
    doc = Document()

    # 添加题目文字
    doc.add_paragraph(text)

    # 添加图片
    image_stream = BytesIO()
    image.save(image_stream, format='PNG')
    image_stream.seek(0)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(image_stream, width=Inches(5.0))

    # 保存文档
    doc.save(output_path)
    print(f"文档已保存到: {output_path}")


def main():
    print("=== SVG-in-JSON 格式测试 ===\n")

    # 1. 解析 JSON
    print("1. 解析 JSON...")
    text, svg_content = parse_svg_json(EXAMPLE_OUTPUT)

    if not svg_content:
        print("错误: 未找到 SVG 内容")
        return

    print(f"✓ 成功解析")
    print(f"  - 文本长度: {len(text)} 字符")
    print(f"  - SVG 长度: {len(svg_content)} 字符")
    print(f"\n题目文字:\n{text[:100]}...\n")

    # 2. SVG 转 PNG
    print("2. 将 SVG 转换为 PNG...")
    try:
        image = svg_to_png(svg_content, output_width=800)
        print(f"✓ 转换成功")
        print(f"  - 图像尺寸: {image.size[0]}x{image.size[1]}")
        print(f"  - 图像模式: {image.mode}")

        # 保存 PNG 用于预览
        png_path = "/mnt/vdb/dev/advanceOCR/output/test_svg_render.png"
        image.save(png_path)
        print(f"  - 已保存到: {png_path}")

    except Exception as e:
        print(f"✗ 转换失败: {e}")
        return

    # 3. 生成 Word 文档
    print("\n3. 生成 Word 文档...")
    output_path = "/mnt/vdb/dev/advanceOCR/output/test_svg_document.docx"
    try:
        create_word_document(text, image, output_path)
        print(f"✓ 文档生成成功!")

    except Exception as e:
        print(f"✗ 文档生成失败: {e}")
        return

    print("\n=== 测试完成 ===")
    print(f"请查看生成的文档: {output_path}")


if __name__ == '__main__':
    main()
