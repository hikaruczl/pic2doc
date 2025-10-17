#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
直接从 LLM 返回的 JSON 内容生成 Word 文档
用法: uv run python generate_from_json.py <json_file> [output.docx]
"""

import sys
import json
import logging
from io import BytesIO
from pathlib import Path
from PIL import Image
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

logging.basicConfig(level=logging.INFO, format='%(levelname)s: %(message)s')
logger = logging.getLogger(__name__)


def parse_json_file(json_path: str) -> dict:
    """读取并解析 JSON 文件"""
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    return data


def svg_to_png(svg_content: str, output_width: int = 800) -> Image.Image:
    """将 SVG 转换为 PNG 图像"""
    try:
        import cairosvg
        png_bytes = cairosvg.svg2png(bytestring=svg_content.encode('utf-8'),
                                      output_width=output_width)
        image = Image.open(BytesIO(png_bytes))
        return image
    except ImportError:
        logger.error("cairosvg 未安装，请运行: uv pip install cairosvg")
        raise
    except Exception as e:
        logger.error(f"SVG 转换失败: {e}")
        raise


def create_word_document(text: str, svg_image: Image.Image = None, output_path: str = "output.docx"):
    """创建包含文本和图片的 Word 文档"""
    doc = Document()

    # 添加题目文字
    doc.add_paragraph(text)

    # 如果有图片，添加到文档中
    if svg_image:
        image_stream = BytesIO()
        svg_image.save(image_stream, format='PNG')
        image_stream.seek(0)

        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(image_stream, width=Inches(5.0))
        logger.info("已添加几何图形到文档")

    # 保存文档
    doc.save(output_path)
    logger.info(f"✓ 文档已保存到: {output_path}")


def main():
    if len(sys.argv) < 2:
        print("用法: uv run python generate_from_json.py <json_file> [output.docx]")
        print("\n示例:")
        print("  uv run python generate_from_json.py llm_output.json")
        print("  uv run python generate_from_json.py llm_output.json custom_output.docx")
        sys.exit(1)

    json_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else "output/generated_from_json.docx"

    # 确保输出目录存在
    output_dir = Path(output_path).parent
    output_dir.mkdir(parents=True, exist_ok=True)

    logger.info(f"正在读取 JSON 文件: {json_path}")

    try:
        # 解析 JSON
        data = parse_json_file(json_path)

        text = data.get('text', '')
        figure_svg = data.get('figure_svg', '')

        if not text:
            logger.warning("JSON 中没有找到 'text' 字段")

        logger.info(f"文本长度: {len(text)} 字符")
        logger.info(f"SVG 长度: {len(figure_svg)} 字符")

        # 如果有 SVG，转换为图像
        svg_image = None
        if figure_svg and figure_svg.strip():
            logger.info("正在渲染 SVG...")
            svg_image = svg_to_png(figure_svg, output_width=800)
            logger.info(f"✓ SVG 渲染成功，尺寸: {svg_image.size[0]}x{svg_image.size[1]}")

        # 生成 Word 文档
        logger.info("正在生成 Word 文档...")
        create_word_document(text, svg_image, output_path)

        print(f"\n✓ 成功生成文档: {output_path}")

    except FileNotFoundError:
        logger.error(f"文件不存在: {json_path}")
        sys.exit(1)
    except json.JSONDecodeError as e:
        logger.error(f"JSON 解析失败: {e}")
        sys.exit(1)
    except Exception as e:
        logger.error(f"生成文档失败: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()
