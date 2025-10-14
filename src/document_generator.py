"""
Word文档生成模块
使用python-docx生成包含公式的Word文档
"""

import os
import re
import logging
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Optional
from io import BytesIO

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image

from .tikz_renderer import TikZRenderer

logger = logging.getLogger(__name__)


class DocumentGenerator:
    """Word文档生成器类"""
    
    def __init__(self, config: dict):
        """
        初始化文档生成器

        Args:
            config: 配置字典
        """
        self.config = config
        self.doc_config = config.get('document', {})
        self.minimal_layout = self.doc_config.get('minimal_layout', False)
        self.output_dir = config.get('output', {}).get('directory', 'output')

        # 确保输出目录存在
        Path(self.output_dir).mkdir(parents=True, exist_ok=True)

        # 初始化TikZ渲染器
        self.tikz_renderer = TikZRenderer(config)

        logger.info("DocumentGenerator initialized (TikZ enabled: %s)",
                   self.tikz_renderer.enabled)
    
    def create_document(self, elements: List[Dict], original_image: Optional[Image.Image] = None,
                       metadata: Optional[Dict] = None, minimal_layout: Optional[bool] = None) -> Document:
        """
        创建Word文档
        
        Args:
            elements: 格式化的元素列表
            original_image: 原始图像 (可选)
            metadata: 元数据 (可选)
            
        Returns:
            Document对象
        """
        doc = Document()

        layout = self.minimal_layout if minimal_layout is None else minimal_layout

        # 始终设置文档属性，便于追踪来源
        self._set_document_properties(doc, metadata)

        if layout:
            # 极简布局：只渲染内容，保持专注于公式预览
            for element in elements:
                self._add_element(doc, element)

            logger.info("文档创建完成（极简布局），共 %s 个元素", len(elements))
            return doc

        # 常规布局
        self._set_page_margins(doc)
        self._add_title(doc, metadata)

        if original_image and self.doc_config.get('include_original_image', True):
            self._add_original_image(doc, original_image)

        doc.add_paragraph('_' * 50)

        for element in elements:
            self._add_element(doc, element)

        self._add_footer(doc, metadata)

        logger.info("文档创建完成,共 %s 个元素", len(elements))
        return doc
    
    def save_document(self, doc: Document, filename: Optional[str] = None) -> str:
        """
        保存文档到文件
        
        Args:
            doc: Document对象
            filename: 文件名 (可选,默认使用时间戳)
            
        Returns:
            保存的文件路径
        """
        if filename is None:
            # 使用配置的文件名模式
            pattern = self.config.get('output', {}).get('filename_pattern', 
                                                        'math_problem_{timestamp}.docx')
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = pattern.format(timestamp=timestamp)
        
        # 确保文件名以.docx结尾
        if not filename.endswith('.docx'):
            filename += '.docx'
        
        filepath = os.path.join(self.output_dir, filename)
        
        try:
            doc.save(filepath)
            logger.info(f"文档已保存: {filepath}")
            return filepath
        except Exception as e:
            logger.error(f"保存文档失败: {str(e)}")
            raise
    
    def _set_document_properties(self, doc: Document, metadata: Optional[Dict]):
        """设置文档属性"""
        core_properties = doc.core_properties
        core_properties.author = "图像转Word系统"
        core_properties.title = "数学问题分析"

        if metadata:
            if 'title' in metadata:
                core_properties.title = metadata['title']
            if 'author' in metadata:
                core_properties.author = metadata['author']
            if 'subject' in metadata:
                core_properties.subject = metadata['subject']
    
    def _set_page_margins(self, doc: Document):
        """设置页边距"""
        sections = doc.sections
        for section in sections:
            margins = self.doc_config.get('page_margins', {})
            section.top_margin = Inches(margins.get('top', 1.0))
            section.bottom_margin = Inches(margins.get('bottom', 1.0))
            section.left_margin = Inches(margins.get('left', 1.0))
            section.right_margin = Inches(margins.get('right', 1.0))
    
    def _add_title(self, doc: Document, metadata: Optional[Dict]):
        """添加标题"""
        title = "数学问题分析结果"
        if metadata and 'title' in metadata:
            title = metadata['title']
        
        heading = doc.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 设置标题字体
        run = heading.runs[0]
        run.font.size = Pt(self.doc_config.get('heading_font_size', 14))
        run.font.name = self.doc_config.get('default_font', 'Arial')
        run.font.color.rgb = RGBColor(0, 0, 0)
    
    def _add_original_image(self, doc: Document, image: Image.Image):
        """添加原始图像"""
        # 添加小标题
        doc.add_heading('原始图像', level=2)
        
        # 保存图像到BytesIO
        img_buffer = BytesIO()
        image.save(img_buffer, format='PNG')
        img_buffer.seek(0)
        
        # 添加图像到文档
        width = self.doc_config.get('image_width_inches', 6.0)
        doc.add_picture(img_buffer, width=Inches(width))
        
        # 添加空行
        doc.add_paragraph()
    
    def _add_element(self, doc: Document, element: Dict):
        """添加单个元素到文档"""
        logger.info("Adding element type=%s content_snippet=%s", element['type'], str(element.get('content', ''))[:100])
        if element['type'] == 'paragraph':
            self._add_paragraph(doc, element['content'])
        
        elif element['type'] == 'text':
            # 处理包含行内公式的文本
            self._add_text_with_inline_formulas(doc, element['content'])
        
        elif element['type'] == 'formula':
            self._add_formula(doc, element)

    def _append_matrix_with_brackets(self, mtable, omml_parent, open_bracket: str, close_bracket: str) -> bool:
        """将MathML的mtable结构转换为带伸缩括号的OMML矩阵"""
        from lxml import etree

        MATH_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'

        matrix_rows: List[List] = []
        max_cols = 0

        for row in mtable:
            row_tag = row.tag.split('}')[-1] if '}' in row.tag else row.tag
            if row_tag != 'mtr':
                continue

            row_cells = []
            for cell in row:
                cell_tag = cell.tag.split('}')[-1] if '}' in cell.tag else cell.tag
                if cell_tag == 'mtd':
                    row_cells.append(cell)

            if row_cells:
                matrix_rows.append(row_cells)
                if len(row_cells) > max_cols:
                    max_cols = len(row_cells)

        if not matrix_rows:
            logger.debug("mtable未找到有效的矩阵行，跳过矩阵转换")
            return False

        d = etree.SubElement(omml_parent, '{%s}d' % MATH_NS)
        dPr = etree.SubElement(d, '{%s}dPr' % MATH_NS)

        begChr = etree.SubElement(dPr, '{%s}begChr' % MATH_NS)
        begChr.set('{%s}val' % MATH_NS, open_bracket)

        endChr = etree.SubElement(dPr, '{%s}endChr' % MATH_NS)
        endChr.set('{%s}val' % MATH_NS, close_bracket)

        matrix_container = etree.SubElement(d, '{%s}e' % MATH_NS)
        m = etree.SubElement(matrix_container, '{%s}m' % MATH_NS)

        if max_cols:
            mPr = etree.SubElement(m, '{%s}mPr' % MATH_NS)
            mcs = etree.SubElement(mPr, '{%s}mcs' % MATH_NS)
            mc = etree.SubElement(mcs, '{%s}mc' % MATH_NS)
            mcPr = etree.SubElement(mc, '{%s}mcPr' % MATH_NS)

            count = etree.SubElement(mcPr, '{%s}count' % MATH_NS)
            count.set('{%s}val' % MATH_NS, str(max_cols))

            mcJc = etree.SubElement(mcPr, '{%s}mcJc' % MATH_NS)
            mcJc.set('{%s}val' % MATH_NS, 'center')

        for row_cells in matrix_rows:
            mr = etree.SubElement(m, '{%s}mr' % MATH_NS)
            for cell in row_cells:
                e = etree.SubElement(mr, '{%s}e' % MATH_NS)

                has_child = False
                for grandchild in cell:
                    has_child = True
                    self._convert_mathml_element_to_omml(grandchild, e)

                if not has_child:
                    text = (cell.text or '').strip()
                    if text:
                        r = etree.SubElement(e, '{%s}r' % MATH_NS)
                        t = etree.SubElement(r, '{%s}t' % MATH_NS)
                        t.text = text

        return True
    
    @staticmethod
    def _clean_xml_incompatible_chars(text: str) -> str:
        """清理XML不兼容的字符（NULL字节和控制字符）"""
        import re
        # 移除NULL字节和XML不兼容的控制字符
        # 保留制表符(\t)、换行符(\n)、回车符(\r)
        # 移除其他控制字符 (0x00-0x08, 0x0B-0x0C, 0x0E-0x1F, 0x7F-0x9F)
        cleaned = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F-\x9F]', '', text)
        return cleaned

    def _add_paragraph(self, doc: Document, text: str):
        """添加段落，支持行内公式和TikZ图形"""
        import re
        from latex2mathml.converter import convert as latex_to_mathml

        # 清理XML不兼容的字符
        text = self._clean_xml_incompatible_chars(text)

        # 扫描所有代码块 (```lang ... ```)
        code_block_pattern = re.compile(r"```(?P<lang>[a-zA-Z0-9_-]+)?\s*\n?(?P<code>.*?)```", re.DOTALL)
        matches = list(code_block_pattern.finditer(text))

        if matches:
            logger.info("Detected %s code block(s) inside paragraph: %s", len(matches), text[:120])
            current_pos = 0

            for block in matches:
                # 处理代码块之前的文本
                if block.start() > current_pos:
                    pre_text = text[current_pos:block.start()].strip()
                    if pre_text:
                        self._add_text_with_inline_formulas(doc, pre_text)

                lang = (block.group('lang') or '').lower()
                code = block.group('code').strip()

                if lang == 'json':
                    self._add_json_block(doc, code)
                    current_pos = block.end()
                    continue

                tikz_code = self._extract_tikz_code(code)
                if tikz_code:
                    logger.info("Rendering TikZ block extracted from %s code block", lang or 'plain')
                    self._add_tikz_figure(doc, tikz_code)
                else:
                    logger.debug("Code block not recognized as TikZ")
                    if lang in {'latex', 'tex'}:
                        cleaned = code.replace('\\(', '$').replace('\\)', '$')
                        cleaned = cleaned.replace('\\[', '$$').replace('\\]', '$$')
                        cleaned = cleaned.replace("\\\\", "\\")
                        paragraphs = re.split(r'\n\s*\n', cleaned)
                        for paragraph_text in paragraphs:
                            paragraph_text = paragraph_text.strip()
                            if paragraph_text:
                                self._add_text_with_inline_formulas(doc, paragraph_text)
                    else:
                        paragraph = doc.add_paragraph()
                        run = paragraph.add_run(code)
                        run.font.name = 'Courier New'
                        run.font.size = Pt(self.doc_config.get('default_font_size', 11))

                current_pos = block.end()

            # 处理剩余文本
            if current_pos < len(text):
                remaining = text[current_pos:].strip()
                if remaining:
                    self._add_text_with_inline_formulas(doc, remaining)
            return

        # 未检测到代码块，检查是否直接包含TikZ环境
        tikz_inline = self._extract_tikz_code(text)
        if tikz_inline:
            logger.info("Rendering inline TikZ environment from paragraph")
            self._add_tikz_figure(doc, tikz_inline)
            # 移除代码块及其围栏
            residual = re.sub(r'```[a-zA-Z0-9_-]*\s*\n?.*?```', '', text, flags=re.DOTALL)
            residual = residual.replace(tikz_inline, '').strip()
            if residual:
                self._add_text_with_inline_formulas(doc, residual)
            return

        # 默认：处理常规文本内容
        self._add_text_with_inline_formulas(doc, text)

    def _add_text_with_inline_formulas(self, doc: Document, text: str):
        """添加文本段落，处理行内公式"""
        import re
        from latex2mathml.converter import convert as latex_to_mathml

        # 统一将 \(\) 与 \[\] 转换为 $ / $$ 以便正则匹配
        text = text.replace('\\(', '$').replace('\\)', '$')
        text = text.replace('\\[', '$$').replace('\\]', '$$')

        paragraph = doc.add_paragraph()

        # 查找所有行内公式 $...$（排除 $$...$$）
        # 使用负向预测和回顾来排除双$
        inline_formula_pattern = r'(?<!\$)\$(?!\$)(.+?)(?<!\$)\$(?!\$)'
        last_pos = 0
        
        for match in re.finditer(inline_formula_pattern, text):
            # 添加公式前的文本
            if match.start() > last_pos:
                run = paragraph.add_run(text[last_pos:match.start()])
                run.font.size = Pt(self.doc_config.get('default_font_size', 11))
                run.font.name = self.doc_config.get('default_font', 'Arial')
            
            # 添加行内公式
            latex = (match.group(1) or match.group(2)).strip()  # 支持$...$或\(...\)
            logger.debug("Inline formula detected: %s", latex)
            latex = self._normalize_inline_latex(latex)
            try:
                mathml = latex_to_mathml(latex)
                self._insert_mathml(paragraph, mathml)
            except Exception as e:
                logger.warning(f"插入行内公式失败，使用文本: {str(e)}")
                # 回退到文本显示
                run = paragraph.add_run(f"${latex}$")
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
            
            last_pos = match.end()
        
        # 添加剩余文本
        if last_pos < len(text):
            run = paragraph.add_run(text[last_pos:])
            run.font.size = Pt(self.doc_config.get('default_font_size', 11))
            run.font.name = self.doc_config.get('default_font', 'Arial')
        
        # 如果段落为空（没有内容），添加文本以避免空段落
        if not paragraph.runs:
            run = paragraph.add_run(text)
            run.font.size = Pt(self.doc_config.get('default_font_size', 11))
            run.font.name = self.doc_config.get('default_font', 'Arial')

    def _add_tikz_figure(self, doc: Document, tikz_code: str):
        """渲染并添加TikZ图形到文档"""
        if not self.tikz_renderer.enabled:
            logger.warning("TikZ renderer is disabled, skipping figure")
            # 添加提示文本
            paragraph = doc.add_paragraph()
            run = paragraph.add_run("[TikZ 图形未渲染：缺少依赖]")
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255, 0, 0)
            return

        try:
            logger.info("Rendering TikZ figure...")
            # 渲染TikZ图形
            image = self.tikz_renderer.render_tikz_to_image(tikz_code)

            if image:
                # 保存图片到BytesIO
                img_buffer = BytesIO()
                image.save(img_buffer, format='PNG')
                img_buffer.seek(0)

                # 添加图片到文档
                width = self.doc_config.get('image_width_inches', 6.0)
                doc.add_picture(img_buffer, width=Inches(width))

                # 添加空行
                doc.add_paragraph()

                logger.info("TikZ figure added successfully")
            else:
                logger.error("TikZ rendering failed, adding placeholder")
                # 添加错误提示
                paragraph = doc.add_paragraph()
                run = paragraph.add_run("[TikZ 图形渲染失败]")
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 0, 0)

        except Exception as e:
            logger.error(f"Error adding TikZ figure: {str(e)}")
            import traceback
            logger.debug(traceback.format_exc())

            # 添加错误提示
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(f"[TikZ 图形渲染错误: {str(e)}]")
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255, 0, 0)

    def _add_json_block(self, doc: Document, code: str):
        """以等宽字体插入JSON文本"""
        lines = code.split('\n')
        for idx, line in enumerate(lines):
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(line)
            run.font.name = 'Courier New'
            run.font.size = Pt(self.doc_config.get('default_font_size', 11))
            if idx == 0:
                paragraph.paragraph_format.space_before = Pt(6)

    @staticmethod
    def _extract_tikz_code(code: str) -> Optional[str]:
        """从代码块中提取TikZ环境"""
        import re

        match = re.search(r'\\begin\{tikzpicture\}.*?\\end\{tikzpicture\}', code, re.DOTALL)
        if match:
            tikz_code = match.group(0).strip()
            logger.debug(f"Extracted TikZ code (first 100 chars): {tikz_code[:100]}...")
            return tikz_code
        return None

    @staticmethod
    def _normalize_inline_latex(latex: str) -> str:
        """将形如 f' 或 f'' 等用 prime 表示的记号转换为标准LaTeX"""
        import re

        def repl(match: re.Match) -> str:
            base = match.group(1)
            primes = match.group(2)
            # 生成正确的上标格式: f^{\prime} 或 f^{\prime\prime}
            prime_str = "\\prime" * len(primes)
            return f"{base}^{{{prime_str}}}"

        # 匹配字母后跟1-4个单引号，但要确保单引号后不是LaTeX命令的一部分
        # 使用负向前瞻避免匹配 f'\text{...} 这样的情况
        converted = re.sub(r"([A-Za-z])('{1,4})(?![a-zA-Z])", repl, latex)
        return converted

    def _add_formula(self, doc: Document, element: Dict):
        """添加公式"""
        formula_type = element['formula_type']
        latex = element['latex']
        mathml = element['mathml']

        # 清理XML不兼容的字符
        latex = self._clean_xml_incompatible_chars(latex)
        mathml = self._clean_xml_incompatible_chars(mathml)

        # 创建段落
        paragraph = doc.add_paragraph()
        paragraph.style = 'Normal'

        # 显示公式的日志
        if formula_type == 'display':
            logger.info(f"准备插入显示公式: {latex[:50]}...")

        # 尝试插入MathML
        try:
            self._insert_mathml(paragraph, mathml, is_display=(formula_type == 'display'))
        except Exception as e:
            logger.warning(f"插入MathML失败,使用LaTeX文本: {str(e)}")
            logger.debug(f"LaTeX: {latex[:100]}...")
            logger.debug(f"MathML: {mathml[:200]}...")
            # 后备方案:插入LaTeX文本
            run = paragraph.add_run(f"${latex}$" if formula_type == 'inline' else f"$${latex}$$")
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
    
    def _insert_mathml(self, paragraph, mathml: str, is_display: bool = False):
        """
        插入MathML到段落

        Args:
            paragraph: 段落对象
            mathml: MathML字符串
            is_display: 是否是显示公式
        """
        from lxml import etree

        # 转换MathML为OMML
        omml_element = self._convert_mathml_to_omml(mathml)

        if omml_element is not None:
            if is_display:
                # 显示公式需要使用oMathPara包装，直接添加到段落元素中（不是run）
                # 这是Word正确显示独立公式的必需结构
                MATH_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'

                # 创建oMathPara元素（display math容器）
                omath_para = etree.Element('{%s}oMathPara' % MATH_NS, nsmap={'m': MATH_NS})

                # 创建oMathParaPr（display math段落属性）
                omath_para_pr = etree.SubElement(omath_para, '{%s}oMathParaPr' % MATH_NS)

                # 添加左对齐设置到oMathParaPr
                jc = etree.SubElement(omath_para_pr, '{%s}jc' % MATH_NS)
                jc.set('{%s}val' % MATH_NS, 'left')

                # 将oMath元素添加到oMathPara中
                omath_para.append(omml_element)

                # 关键：直接添加oMathPara到段落元素（不通过run）
                # 这是display math和inline math的关键区别
                paragraph._element.append(omath_para)

                logger.info(f"显示公式已插入（使用oMathPara左对齐）")
            else:
                # 行内公式：插入到run中（这是正确的inline math结构）
                run = paragraph.add_run()
                run._element.append(omml_element)

                logger.debug(f"行内公式已插入")
        else:
            raise Exception("MathML to OMML conversion failed")
    
    def _convert_mathml_to_omml(self, mathml: str):
        """
        将MathML转换为OMML (Office Math Markup Language)
        
        Args:
            mathml: MathML字符串
            
        Returns:
            OMML lxml Element对象,如果转换失败返回None
        """
        try:
            from lxml import etree
            
            # 解析MathML
            mathml_clean = mathml.strip()
            if not mathml_clean.startswith('<math'):
                logger.debug(f"MathML格式错误: 不以<math开头")
                return None
            
            # 解析MathML
            try:
                root = etree.fromstring(mathml_clean.encode('utf-8'))
            except Exception as e:
                logger.debug(f"MathML解析失败: {str(e)}")
                return None
            
            # 创建OMML根元素
            MATH_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
            omath = etree.Element('{%s}oMath' % MATH_NS, nsmap={'m': MATH_NS})
            
            # 递归转换MathML元素为OMML
            self._convert_mathml_element_to_omml(root, omath)
            
            logger.debug(f"OMML转换成功: {len(list(omath))} 个子元素")
            return omath
            
        except Exception as e:
            logger.debug(f"MathML到OMML转换失败: {str(e)}")
            import traceback
            logger.debug(traceback.format_exc())
            return None
    
    def _convert_mathml_element_to_omml(self, mathml_elem, omml_parent):
        """
        递归转换MathML元素为OMML元素

        Args:
            mathml_elem: MathML lxml元素
            omml_parent: OMML lxml父元素
        """
        from lxml import etree

        MATH_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'

        # 获取标签名(去除命名空间)
        tag = mathml_elem.tag.split('}')[-1] if '}' in mathml_elem.tag else mathml_elem.tag

        # 处理不同的MathML标签
        if tag == 'math':
            # 根元素,处理子元素
            for child in mathml_elem:
                self._convert_mathml_element_to_omml(child, omml_parent)

        elif tag == 'mrow':
            children = list(mathml_elem)
            total_children = len(children)
            i = 0

            while i < total_children:
                child = children[i]
                child_tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

                # 检测模式: mo + mtable + mo (允许中间存在mspace)
                if child_tag == 'mo':
                    open_bracket = (child.text or '(').strip() or '('

                    # 跳过空白节点
                    mid_index = i + 1
                    while mid_index < total_children:
                        mid_child = children[mid_index]
                        mid_tag = mid_child.tag.split('}')[-1] if '}' in mid_child.tag else mid_child.tag
                        if mid_tag == 'mspace':
                            mid_index += 1
                            continue
                        break
                    else:
                        mid_child = None
                        mid_tag = None

                    if mid_child is not None and mid_tag == 'mtable':
                        end_index = mid_index + 1
                        while end_index < total_children:
                            end_child = children[end_index]
                            end_tag = end_child.tag.split('}')[-1] if '}' in end_child.tag else end_child.tag
                            if end_tag == 'mspace':
                                end_index += 1
                                continue
                            break
                        else:
                            end_child = None
                            end_tag = None

                        if end_child is not None and end_tag == 'mo':
                            close_bracket = (end_child.text or ')').strip() or ')'

                            converted = self._append_matrix_with_brackets(
                                mid_child, omml_parent, open_bracket, close_bracket
                            )

                            if converted:
                                i = end_index + 1
                                continue

                # 默认处理
                self._convert_mathml_element_to_omml(child, omml_parent)
                i += 1

        elif tag in ['mi', 'mn', 'mo', 'mtext']:
            # 标识符、数字、运算符、文本
            r = etree.SubElement(omml_parent, '{%s}r' % MATH_NS)
            t = etree.SubElement(r, '{%s}t' % MATH_NS)
            text = mathml_elem.text or ''
            t.text = text

        elif tag == 'mfrac':
            # 分数
            f = etree.SubElement(omml_parent, '{%s}f' % MATH_NS)
            num = etree.SubElement(f, '{%s}num' % MATH_NS)
            den = etree.SubElement(f, '{%s}den' % MATH_NS)

            children = list(mathml_elem)
            if len(children) >= 2:
                self._convert_mathml_element_to_omml(children[0], num)
                self._convert_mathml_element_to_omml(children[1], den)

        elif tag == 'msup':
            # 上标
            ssup = etree.SubElement(omml_parent, '{%s}sSup' % MATH_NS)
            e = etree.SubElement(ssup, '{%s}e' % MATH_NS)
            sup = etree.SubElement(ssup, '{%s}sup' % MATH_NS)

            children = list(mathml_elem)
            if len(children) >= 2:
                self._convert_mathml_element_to_omml(children[0], e)
                self._convert_mathml_element_to_omml(children[1], sup)

        elif tag == 'msub':
            # 下标
            ssub = etree.SubElement(omml_parent, '{%s}sSub' % MATH_NS)
            e = etree.SubElement(ssub, '{%s}e' % MATH_NS)
            sub = etree.SubElement(ssub, '{%s}sub' % MATH_NS)

            children = list(mathml_elem)
            if len(children) >= 2:
                self._convert_mathml_element_to_omml(children[0], e)
                self._convert_mathml_element_to_omml(children[1], sub)

        elif tag == 'msubsup':
            # 同时有下标和上标（问题1的修复）
            ssubsup = etree.SubElement(omml_parent, '{%s}sSubSup' % MATH_NS)
            e = etree.SubElement(ssubsup, '{%s}e' % MATH_NS)
            sub = etree.SubElement(ssubsup, '{%s}sub' % MATH_NS)
            sup = etree.SubElement(ssubsup, '{%s}sup' % MATH_NS)

            children = list(mathml_elem)
            if len(children) >= 3:
                self._convert_mathml_element_to_omml(children[0], e)
                self._convert_mathml_element_to_omml(children[1], sub)
                self._convert_mathml_element_to_omml(children[2], sup)

        elif tag == 'msqrt':
            # 平方根
            rad = etree.SubElement(omml_parent, '{%s}rad' % MATH_NS)
            radPr = etree.SubElement(rad, '{%s}radPr' % MATH_NS)
            degHide = etree.SubElement(radPr, '{%s}degHide' % MATH_NS)
            degHide.set('{%s}val' % MATH_NS, '1')
            e = etree.SubElement(rad, '{%s}e' % MATH_NS)

            for child in mathml_elem:
                self._convert_mathml_element_to_omml(child, e)

        elif tag == 'mroot':
            # n次根
            rad = etree.SubElement(omml_parent, '{%s}rad' % MATH_NS)
            deg = etree.SubElement(rad, '{%s}deg' % MATH_NS)
            e = etree.SubElement(rad, '{%s}e' % MATH_NS)

            children = list(mathml_elem)
            if len(children) >= 2:
                self._convert_mathml_element_to_omml(children[0], e)
                self._convert_mathml_element_to_omml(children[1], deg)

        elif tag == 'mover':
            # 上方符号（问题2的修复：向量、上划线等）
            # 检查是否是accent（重音符号）
            accent_attr = mathml_elem.get('accent', 'false')

            if accent_attr == 'true':
                # 使用accent结构
                acc = etree.SubElement(omml_parent, '{%s}acc' % MATH_NS)
                accPr = etree.SubElement(acc, '{%s}accPr' % MATH_NS)
                e = etree.SubElement(acc, '{%s}e' % MATH_NS)

                children = list(mathml_elem)
                if len(children) >= 2:
                    # 第一个子元素是基础，第二个是上方符号
                    self._convert_mathml_element_to_omml(children[0], e)
                    # 设置accent字符
                    chr_elem = etree.SubElement(accPr, '{%s}chr' % MATH_NS)
                    accent_text = children[1].text if children[1].text else '→'
                    chr_elem.set('{%s}val' % MATH_NS, accent_text)
            else:
                # 使用bar结构（上划线）或limUpp结构
                # 检查上方是否是overline或bar
                children = list(mathml_elem)
                if len(children) >= 2:
                    over_elem = children[1]
                    over_text = over_elem.text if over_elem.text else ''

                    # 检查是否是上划线
                    if '‾' in over_text or 'OverBar' in over_text or over_text == '¯':
                        # 使用bar结构
                        bar = etree.SubElement(omml_parent, '{%s}bar' % MATH_NS)
                        barPr = etree.SubElement(bar, '{%s}barPr' % MATH_NS)
                        pos = etree.SubElement(barPr, '{%s}pos' % MATH_NS)
                        pos.set('{%s}val' % MATH_NS, 'top')
                        e = etree.SubElement(bar, '{%s}e' % MATH_NS)
                        self._convert_mathml_element_to_omml(children[0], e)
                    else:
                        # 使用limUpp结构（上方符号，如向量箭头）
                        limUpp = etree.SubElement(omml_parent, '{%s}limUpp' % MATH_NS)
                        e = etree.SubElement(limUpp, '{%s}e' % MATH_NS)
                        lim = etree.SubElement(limUpp, '{%s}lim' % MATH_NS)

                        self._convert_mathml_element_to_omml(children[0], e)
                        self._convert_mathml_element_to_omml(children[1], lim)

        elif tag == 'munder':
            # 下方符号（下划线等）
            # 使用limLow结构或bar结构
            children = list(mathml_elem)
            if len(children) >= 2:
                under_elem = children[1]
                under_text = under_elem.text if under_elem.text else ''

                # 检查是否是下划线
                if '_' in under_text or 'UnderBar' in under_text:
                    # 使用bar结构
                    bar = etree.SubElement(omml_parent, '{%s}bar' % MATH_NS)
                    barPr = etree.SubElement(bar, '{%s}barPr' % MATH_NS)
                    pos = etree.SubElement(barPr, '{%s}pos' % MATH_NS)
                    pos.set('{%s}val' % MATH_NS, 'bot')
                    e = etree.SubElement(bar, '{%s}e' % MATH_NS)
                    self._convert_mathml_element_to_omml(children[0], e)
                else:
                    # 使用limLow结构
                    limLow = etree.SubElement(omml_parent, '{%s}limLow' % MATH_NS)
                    e = etree.SubElement(limLow, '{%s}e' % MATH_NS)
                    lim = etree.SubElement(limLow, '{%s}lim' % MATH_NS)

                    self._convert_mathml_element_to_omml(children[0], e)
                    self._convert_mathml_element_to_omml(children[1], lim)

        elif tag == 'mtable':
            # 矩阵表格 - 检查父元素是否包含括号
            parent = mathml_elem.getparent()
            has_brackets = False

            if parent is not None:
                parent_tag = parent.tag.split('}')[-1] if '}' in parent.tag else parent.tag
                if parent_tag == 'mrow':
                    # 检查mrow的第一个和最后一个子元素是否是括号
                    siblings = list(parent)
                    if len(siblings) >= 3:
                        first_tag = siblings[0].tag.split('}')[-1] if '}' in siblings[0].tag else siblings[0].tag
                        last_tag = siblings[-1].tag.split('}')[-1] if '}' in siblings[-1].tag else siblings[-1].tag
                        if first_tag == 'mo' and last_tag == 'mo':
                            has_brackets = True

            if has_brackets:
                # 矩阵已经在mrow处理时添加了括号，这里只处理内容
                # 跳过，让mrow处理
                return
            else:
                # 表格结构(用于aligned/gathered环境或无括号矩阵)
                # OMML中使用eqArr(方程组数组)来表示
                eqArr = etree.SubElement(omml_parent, '{%s}eqArr' % MATH_NS)

                # 处理每一行
                for row in mathml_elem:
                    row_tag = row.tag.split('}')[-1] if '}' in row.tag else row.tag
                    if row_tag == 'mtr':  # 表格行
                        e = etree.SubElement(eqArr, '{%s}e' % MATH_NS)

                        # 处理行中的单元格
                        for cell in row:
                            cell_tag = cell.tag.split('}')[-1] if '}' in cell.tag else cell.tag
                        if cell_tag == 'mtd':  # 表格单元格
                            for child in cell:
                                self._convert_mathml_element_to_omml(child, e)

        elif tag == 'mtr':
            # 表格行(如果直接遇到,不在mtable中)
            for child in mathml_elem:
                self._convert_mathml_element_to_omml(child, omml_parent)

        elif tag == 'mtd':
            # 表格单元格(如果直接遇到,不在mtr中)
            for child in mathml_elem:
                self._convert_mathml_element_to_omml(child, omml_parent)

        elif tag == 'mfenced':
            # 括号结构（问题3的修复：矩阵括号）
            # 获取开闭括号属性
            open_bracket = mathml_elem.get('open', '(')
            close_bracket = mathml_elem.get('close', ')')

            # 检查是否包含mtable（矩阵）
            has_mtable = False
            for child in mathml_elem:
                child_tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                if child_tag == 'mtable':
                    has_mtable = True
                    break

            if has_mtable:
                # 矩阵情况：使用m结构
                d = etree.SubElement(omml_parent, '{%s}d' % MATH_NS)
                dPr = etree.SubElement(d, '{%s}dPr' % MATH_NS)

                # 设置括号字符
                begChr = etree.SubElement(dPr, '{%s}begChr' % MATH_NS)
                begChr.set('{%s}val' % MATH_NS, open_bracket)
                endChr = etree.SubElement(dPr, '{%s}endChr' % MATH_NS)
                endChr.set('{%s}val' % MATH_NS, close_bracket)

                # 处理矩阵内容
                for child in mathml_elem:
                    child_tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                    if child_tag == 'mtable':
                        # 创建矩阵结构
                        m = etree.SubElement(d, '{%s}m' % MATH_NS)

                        # 处理每一行
                        for row in child:
                            row_tag = row.tag.split('}')[-1] if '}' in row.tag else row.tag
                            if row_tag == 'mtr':
                                mr = etree.SubElement(m, '{%s}mr' % MATH_NS)
                                # 处理每个单元格
                                for cell in row:
                                    cell_tag = cell.tag.split('}')[-1] if '}' in cell.tag else cell.tag
                                    if cell_tag == 'mtd':
                                        e = etree.SubElement(mr, '{%s}e' % MATH_NS)
                                        for grandchild in cell:
                                            self._convert_mathml_element_to_omml(grandchild, e)
            else:
                # 普通括号：使用d结构
                d = etree.SubElement(omml_parent, '{%s}d' % MATH_NS)
                dPr = etree.SubElement(d, '{%s}dPr' % MATH_NS)

                # 设置括号字符
                begChr = etree.SubElement(dPr, '{%s}begChr' % MATH_NS)
                begChr.set('{%s}val' % MATH_NS, open_bracket)
                endChr = etree.SubElement(dPr, '{%s}endChr' % MATH_NS)
                endChr.set('{%s}val' % MATH_NS, close_bracket)

                # 处理内容
                e = etree.SubElement(d, '{%s}e' % MATH_NS)
                for child in mathml_elem:
                    self._convert_mathml_element_to_omml(child, e)

        elif tag == 'mspace':
            # 空格 - 在OMML中添加空文本
            r = etree.SubElement(omml_parent, '{%s}r' % MATH_NS)
            t = etree.SubElement(r, '{%s}t' % MATH_NS)
            t.text = ' '

        else:
            # 未处理的标签,尝试处理子元素
            logger.debug(f"未处理的MathML标签: {tag}")
            for child in mathml_elem:
                self._convert_mathml_element_to_omml(child, omml_parent)
    
    def _add_footer(self, doc: Document, metadata: Optional[Dict]):
        """添加页脚"""
        # 添加分隔线
        doc.add_paragraph('_' * 50)

        # 添加生成信息
        footer_text = f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"

        footer = doc.add_paragraph(footer_text)
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = footer.runs[0]
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(128, 128, 128)
        run.italic = True
    
    def create_from_analysis(self, analysis_result: Dict, original_image: Image.Image,
                            elements: List[Dict], filename: Optional[str] = None,
                            minimal_layout: Optional[bool] = None) -> str:
        """
        从分析结果创建并保存文档

        Args:
            analysis_result: LLM分析结果
            original_image: 原始图像
            elements: 格式化的元素列表
            filename: 可选的输出文件名

        Returns:
            保存的文件路径
        """
        # 准备元数据
        metadata = {
            'title': '数学问题分析',
            'provider': f"{analysis_result.get('provider', 'unknown')} - {analysis_result.get('model', 'unknown')}",
            'subject': 'Math Problem OCR'
        }

        # 创建文档
        doc = self.create_document(elements, original_image, metadata, minimal_layout=minimal_layout)

        # 保存文档
        filepath = self.save_document(doc, filename)

        return filepath
