# -*- coding: utf-8 -*-
"""
测试公式插入修复
"""
from pathlib import Path
import sys
import yaml

ROOT_DIR = Path(__file__).resolve().parents[2]
if str(ROOT_DIR) not in sys.path:
    sys.path.insert(0, str(ROOT_DIR))

from src.document_generator import DocumentGenerator
from src.formula_converter import FormulaConverter
from docx import Document

# 加载配置
with open(ROOT_DIR / 'config' / 'config.yaml', 'r', encoding='utf-8') as f:
    config = yaml.safe_load(f)

# 创建实例
doc_gen = DocumentGenerator(config)
formula_converter = FormulaConverter(config)

# 测试简单公式
test_formulas = [
    ("E = mc^2", "inline"),
    ("x^2 + y^2 = z^2", "inline"),
    ("\\frac{a}{b}", "display"),
    ("\\sqrt{x}", "inline"),
]

print("测试公式转换...")
print("=" * 60)

for latex, formula_type in test_formulas:
    print(f"\n测试: {latex} ({formula_type})")
    
    # 转换为MathML
    mathml = formula_converter.convert_latex_to_mathml(latex)
    print(f"MathML: {mathml[:100]}...")
    
    # 测试OMML转换
    try:
        omml = doc_gen._convert_mathml_to_omml(mathml)
        if omml:
            print(f"✓ OMML转换成功 (长度: {len(omml)})")
            print(f"OMML预览: {omml[:200]}...")
        else:
            print("✗ OMML转换失败")
    except Exception as e:
        print(f"✗ 转换出错: {str(e)}")

# 创建测试文档
print("\n" + "=" * 60)
print("创建测试文档...")

doc = Document()
doc.add_heading('公式测试文档', level=1)

for latex, formula_type in test_formulas:
    doc.add_paragraph(f"LaTeX: {latex}")
    
    # 创建公式元素
    mathml = formula_converter.convert_latex_to_mathml(latex)
    
    element = {
        'type': 'formula',
        'formula_type': formula_type,
        'latex': latex,
        'mathml': mathml
    }
    
    try:
        doc_gen._add_formula(doc, element)
        print(f"✓ 添加公式成功: {latex}")
    except Exception as e:
        print(f"✗ 添加公式失败: {latex} - {str(e)}")

# 保存测试文档
output_path = ROOT_DIR / 'output' / 'test_formula_fix.docx'
try:
    doc.save(output_path)
    print(f"\n✓ 测试文档已保存: {output_path}")
    print("\n请在Word中打开文档检查公式是否可编辑")
except Exception as e:
    print(f"\n✗ 保存文档失败: {str(e)}")
